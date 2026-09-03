from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

B=Path(r'C:\ai_web\lecture1\인테리어사이트\outputs\거실-인테리어-제안서')
P=B/'01_현장사진.png'; I=B/'02_모던클래식_디자인시안.png'; O=B/'공간한쪽_거실_모던클래식_디자인제안서.docx'
BLUE='2864B4'; NAVY='173B67'; INK='28374A'; MUTED='667085'; PALE='EAF2FC'; WARM='F5F0E8'; FONT='Malgun Gothic'
d=Document(); s=d.sections[0]; s.page_width=Inches(8.5); s.page_height=Inches(11); s.top_margin=s.bottom_margin=Inches(.72); s.left_margin=s.right_margin=Inches(.82)
def f(r,z=10.3,b=False,c=INK,i=False):
 r.font.name=FONT; rp=r._element.get_or_add_rPr(); [rp.rFonts.set(qn('w:'+x),FONT) for x in ('ascii','hAnsi','eastAsia')]; r.font.size=Pt(z); r.bold=b; r.italic=i; r.font.color.rgb=RGBColor.from_string(c)
def p(txt='',z=10.3,b=False,c=INK,al=None,after=7):
 x=d.add_paragraph(); x.paragraph_format.space_after=Pt(after); x.paragraph_format.line_spacing=1.24; x.alignment=al; f(x.add_run(txt),z,b,c); return x
def h(txt,l=1): d.add_heading(txt,level=l)
def bullets(items):
 for x in items:
  q=d.add_paragraph(style='List Bullet'); q.paragraph_format.space_after=Pt(4); f(q.add_run(x),10)
def box(label,txt):
 t=d.add_table(rows=1,cols=1); t.style='Table Grid'; c=t.cell(0,0); sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),PALE); c._tc.get_or_add_tcPr().append(sh); q=c.paragraphs[0]; f(q.add_run(label+'  '),10.4,True,NAVY); f(q.add_run(txt),10.4); p('',after=0)
def page(): d.add_page_break()
for n,z in [('Heading 1',16),('Heading 2',13)]:
 st=d.styles[n]; st.font.name=FONT; st._element.rPr.rFonts.set(qn('w:eastAsia'),FONT); st.font.size=Pt(z); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(BLUE); st.paragraph_format.space_before=Pt(14); st.paragraph_format.space_after=Pt(7); st.paragraph_format.keep_with_next=True
st=d.styles['Normal']; st.font.name=FONT; st._element.rPr.rFonts.set(qn('w:eastAsia'),FONT); st.font.size=Pt(10.3)
head=s.header.paragraphs[0]; head.alignment=WD_ALIGN_PARAGRAPH.RIGHT; f(head.add_run('SPACE HANCHOK · DESIGN REFERENCE'),8.5,True,BLUE)
foot=s.footer.paragraphs[0]; foot.alignment=WD_ALIGN_PARAGRAPH.RIGHT; f(foot.add_run('공간한쪽 | 거실 디자인 제안   '),8.5,False,MUTED); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); foot._p.append(fld)

p('공간한쪽',11,True,BLUE); p('거실 모던 클래식 디자인 제안서',24,True,NAVY,after=3); p('실제 사례에서 원리를 추출하고, 이 거실의 마감으로 다시 설계합니다.',12.3,False,MUTED,after=10)
d.add_picture(str(I),width=Inches(6.85)); p('완성 시안 · 패널 몰딩, 웜그레이지, 올리브그레이 포인트, 레이어드 조명과 이중 커튼',8.3,False,MUTED,WD_ALIGN_PARAGRAPH.CENTER)
box('디자인 우선','가구는 구매 대상이 아니라 스케일과 분위기를 확인하기 위한 연출입니다. 중심은 벽·천장·바닥 경계와 빛입니다.')

page(); h('1. 레퍼런스 분석',1); p('국내 아파트의 낮은 천장고와 시스템에어컨 조건, 해외 모던 클래식의 비례·패널·커튼·조명 사례를 함께 비교했습니다.')
h('공통적으로 추출한 6가지',2); bullets(['장식 몰딩의 개수보다 패널 비례와 정렬이 완성도를 좌우합니다.','벽과 몰딩을 같은 색으로 칠하면 클래식 요소가 과장되지 않습니다.','중성색 뼈대에 올리브·블루그레이 같은 저채도 색을 10~20%만 사용합니다.','창 전체를 감싸는 쉬어·암막 이중 커튼이 딱딱한 몰딩을 부드럽게 합니다.','천장 한 등만 밝히지 않고 장식등·벽등·스탠드로 빛을 겹칩니다.','금속은 브라스 한 종류만 소량 사용하고 대리석·골드 장식은 절제합니다.'])
h('참고 사례 10선',2); bullets(['공감디자인 · 광명 50평대 모던 프렌치 클래식','공감디자인 · 논현동 클래식 몰딩 아파트','Archisketch · 30평대 모던 클래식 거실','Interior Design Seoul · Reinvented Classic Elegance','Livingetc · Architectural Molding Ideas','Homes & Gardens · Living Room Wall Lighting','Martha Stewart · Living Room Curtain Ideas','Livingetc · Designer Curtain Hanging','Wallpaper · Seoul Apartment Interior','Modern Classic Interior Styling reference book'])

page(); h('2. 이 거실에 적용한 디자인',1); d.add_picture(str(P),width=Inches(4.2)); p('현재: 오른쪽 진한 수평 패널이 시선을 분절하고, 주조명이 평면적으로 밝히는 상태',8.3,False,MUTED,WD_ALIGN_PARAGRAPH.CENTER)
h('선택 요소 → 적용 → 이유',2); bullets(['얇은 박스 몰딩 → 오른쪽 기존 패널벽 철거 후 큰 패널 3개와 하부 패널 구성 → 수평선을 없애고 천장고가 높아 보이게 합니다.','웜그레이지 에그쉘 도장 → 벽·몰딩·문선에 연결 → 기존 밝은 바닥과 충돌하지 않고 몰딩 그림자가 살아납니다.','올리브그레이 포인트 → 큰 패널 내부 일부만 적용 → 클래식의 중후함을 주되 공간을 어둡게 만들지 않습니다.','슬림 코니스와 평걸레받이 → 벽과 같은 색으로 도장 → 모던 클래식의 경계를 만들면서 관리성을 유지합니다.','오팔 글로브 조명 → 기존 전원 위치에서 교체 → 시스템에어컨을 막지 않고 중심성을 만듭니다.','쉬어+암막 이중커튼 → 창 전체 높이로 설치 → 외부 아파트 전망을 부드럽게 흐리고 벽면 비례를 완성합니다.'])
box('현장 적용 제외','벽난로, 빅슬랩 아트월, 과도한 골드, 복잡한 천장 몰딩, 무리한 간접천장은 이 거실의 크기와 기존 설비에 비해 과합니다.')

page(); h('3. 마감 공사와 계획금액',1); h('오른쪽 기존 패널벽',2); bullets(['철거·폐기 40~90만원: 접착 방식과 바탕 손상에 따라 변동','면 보수·전체 퍼티 30~80만원: 패널 철거 후 접착제·요철 제거','박스 몰딩 80~160만원: 재질, 총 길이, 코너 수, 도장 포함 여부에 따라 변동'])
h('도장과 벽지',2); bullets(['전체 수성 도장 90~180만원: 벽지 제거, 퍼티, 샌딩, 프라이머, 상도 2회 기준','실크 벽지 60~120만원: 요철을 숨기기 쉽지만 몰딩과 한 면 같은 일체감은 약함','권장: 패널벽과 몰딩은 도장, 몰딩 없는 보조 벽은 예산에 따라 도장 또는 무지 실크벽지'])
h('상·하부 경계',2); bullets(['슬림 코니스+평걸레받이 교체 50~120만원','상부 무몰딩 70~150만원: 철거 자국과 천장 수평 보수 포함','마이너스 몰딩/매립 걸레받이 150~320만원: 목공·석고·미장·도장을 동반해 부분공사 효율이 낮음'])
h('조명·커튼',2); bullets(['오팔 장식등 제품 15~60만원 + 교체 5~15만원','벽등 또는 픽처라이트 1~2개 배선 포함 25~70만원','쉬어+암막 이중커튼 40~90만원, 레일·주름배수·원단에 따라 변동'])
box('예산 범위','권장안은 약 350~700만원입니다. 실측 후 철거, 면 보수, 몰딩 길이, 도장 횟수, 전기, 커튼을 분리 견적해야 합니다.')

page(); h('4. 시공 순서와 확정 기준',1); bullets(['벽별 실측과 패널 전개도 작성','레퍼런스 중 선호/비선호 표시','A3 이상 컬러 샘플을 주간·야간에 확인','기존 패널 철거 후 바탕 상태 재견적','전기 배선과 등기구 보강 선행','몰딩 목공 → 퍼티 → 샌딩 → 프라이머 → 상도 2회','코니스·걸레받이·커튼·등기구 마감','사광에서 몰딩 이음과 벽면 울렁임 최종 검수'])
h('실측 후 확정할 값',2); bullets(['오른쪽 벽 전체 폭·높이와 인터폰·스위치 좌표','기존 패널 재질·접착 방식·철거 후 바탕','패널 몰딩 폭 20~30mm, 돌출 8~12mm, 패널 간격','주조명 무게·직경·에어컨 토출 간격','창 폭·높이·커튼박스 깊이·커튼 완성 폭','벽지 유지 범위와 전체 퍼티 필요 여부'])
box('다음 제안 단계','선호 레퍼런스 3~5개를 고르면 패널 전개 비례, 정확한 색상, 조명 형태를 2안으로 좁혀 L2 시안과 공정별 수량 견적으로 발전시킵니다.')
d.core_properties.title='공간한쪽 거실 모던 클래식 디자인 제안서'; d.core_properties.author='공간한쪽'; d.save(O); print(O)
