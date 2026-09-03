from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

BASE = Path(r"C:\ai_web\lecture1\인테리어사이트\outputs\방-인테리어-제안서")
PHOTO = BASE / "01_현장사진.png"
CONCEPT = BASE / "03_모던클래식_디자인시안.png"
OUT = BASE / "공간한쪽_방_모던클래식_디자인제안서.docx"

BLUE, NAVY, INK, MUTED = "2864B4", "173B67", "28374A", "667085"
PALE, WARM, LINE, GOLD = "EAF2FC", "F5F0E8", "D8E1EC", "8A671A"
FONT = "Malgun Gothic"

doc = Document(); sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(.7)
sec.left_margin = sec.right_margin = Inches(.8)
sec.header_distance = sec.footer_distance = Inches(.35)

def font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    for key in ('eastAsia','ascii','hAnsi'): rpr.rFonts.set(qn('w:'+key), FONT)
    run.font.size, run.bold, run.italic = Pt(size), bold, italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run

normal = doc.styles['Normal']; normal.font.name = FONT; normal.font.size = Pt(10.3)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(7); normal.paragraph_format.line_spacing = 1.25
for n,s,c,b,a in [('Heading 1',16,BLUE,18,9),('Heading 2',13,BLUE,12,6),('Heading 3',11.5,NAVY,8,4)]:
    st=doc.styles[n]; st.font.name=FONT; st._element.rPr.rFonts.set(qn('w:eastAsia'),FONT)
    st.font.size=Pt(s); st.font.bold=True; st.font.color.rgb=RGBColor.from_string(c)
    st.paragraph_format.space_before=Pt(b); st.paragraph_format.space_after=Pt(a); st.paragraph_format.keep_with_next=True
for n in ('List Bullet','List Number'):
    st=doc.styles[n]; st.font.name=FONT; st._element.rPr.rFonts.set(qn('w:eastAsia'),FONT); st.font.size=Pt(10)
    st.paragraph_format.left_indent=Inches(.38); st.paragraph_format.first_line_indent=Inches(-.19)
    st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.18

def shade(cell, fill):
    pr=cell._tc.get_or_add_tcPr(); sh=pr.find(qn('w:shd'))
    if sh is None: sh=OxmlElement('w:shd'); pr.append(sh)
    sh.set(qn('w:fill'),fill)

def margins(cell):
    pr=cell._tc.get_or_add_tcPr(); mar=pr.first_child_found_in('w:tcMar')
    if mar is None: mar=OxmlElement('w:tcMar'); pr.append(mar)
    for n,v in [('top',90),('start',120),('bottom',90),('end',120)]:
        x=mar.find(qn('w:'+n))
        if x is None: x=OxmlElement('w:'+n); mar.append(x)
        x.set(qn('w:w'),str(v)); x.set(qn('w:type'),'dxa')

def geometry(t, widths):
    t.autofit=False; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    total=int(sum(widths)*1440); pr=t._tbl.tblPr
    tw=pr.find(qn('w:tblW'))
    if tw is None: tw=OxmlElement('w:tblW'); pr.append(tw)
    tw.set(qn('w:w'),str(total)); tw.set(qn('w:type'),'dxa')
    ti=pr.find(qn('w:tblInd'))
    if ti is None: ti=OxmlElement('w:tblInd'); pr.append(ti)
    ti.set(qn('w:w'),'120'); ti.set(qn('w:type'),'dxa')
    grid=t._tbl.tblGrid
    for x in list(grid): grid.remove(x)
    for w in widths:
        x=OxmlElement('w:gridCol'); x.set(qn('w:w'),str(int(w*1440))); grid.append(x)
    for row in t.rows:
        for i,c in enumerate(row.cells):
            c.width=Inches(widths[i]); margins(c)
            cw=c._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if cw is None: cw=OxmlElement('w:tcW'); c._tc.get_or_add_tcPr().append(cw)
            cw.set(qn('w:w'),str(int(widths[i]*1440))); cw.set(qn('w:type'),'dxa')

def table(headers, rows, widths):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    for i,h in enumerate(headers): t.rows[0].cells[i].text=h
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    geometry(t,widths)
    for ri,row in enumerate(t.rows):
        for c in row.cells:
            c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ri==0: shade(c,PALE)
            for p in c.paragraphs:
                p.paragraph_format.space_after=Pt(2); p.paragraph_format.line_spacing=1.13
                for r in p.runs: font(r,8.9,ri==0,NAVY if ri==0 else INK)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)
    return t

def callout(label,text,fill=WARM,color=GOLD):
    t=doc.add_table(rows=1,cols=1); t.style='Table Grid'; geometry(t,[6.9]); c=t.cell(0,0); shade(c,fill)
    p=c.paragraphs[0]; p.paragraph_format.space_after=Pt(0)
    font(p.add_run(label+'  '),10.3,True,color); font(p.add_run(text),10.3)
    doc.add_paragraph().paragraph_format.space_after=Pt(0)

def bullet(text): p=doc.add_paragraph(style='List Bullet'); font(p.add_run(text),10); return p
def caption(text):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(7)
    font(p.add_run(text),8.4,False,MUTED,True)
def page(): doc.add_page_break()

head=sec.header.paragraphs[0]; head.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(head.add_run('SPACE HANCHOK · DESIGN FIRST'),8.5,True,BLUE)
foot=sec.footer.paragraphs[0]; foot.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(foot.add_run('공간한쪽 | 모던 클래식 디자인 제안   '),8.5,False,MUTED)
fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); foot._p.append(fld)

# 1 Cover
p=doc.add_paragraph(); font(p.add_run('공간한쪽'),11,True,BLUE)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3); font(p.add_run('모던 클래식 방 디자인 제안서'),24,True,NAVY)
p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(11); font(p.add_run('가구보다 먼저, 벽과 경계와 빛을 설계합니다.'),12.5,False,MUTED)
doc.add_picture(str(CONCEPT),width=Inches(6.9)); caption('디자인 시안 · 벽면 패널 몰딩, 저채도 포인트 컬러, 코니스, 오팔 조명, 이중 커튼')
callout('디자인 우선','가구는 분위기와 스케일을 설명하기 위한 연출입니다. 구매 품목과 동선 최적화는 이번 제안의 중심에서 제외합니다.')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(p.add_run('2026. 09. 02 | 사진 기반 계획견적 L1'),9,False,MUTED)

page(); doc.add_heading('1. 디자인 콘셉트',1)
doc.add_picture(str(PHOTO),width=Inches(3.0)); caption('현재 공간: 밝지만 벽과 천장 경계의 표정이 약한 상태')
callout('핵심 문장','장식은 얇게, 비례는 정확하게, 컬러는 낮게. 클래식 요소를 더하되 오래 보아도 부담 없는 방을 만듭니다.',PALE,NAVY)
doc.add_heading('제안 팔레트',2)
table(['역할','권장 색','적용 위치','주의점'],[
('기본','웜그레이지 #D8D0C4 계열','벽·몰딩 동일색','북향이면 한 톤 밝게'),
('포인트','더스티 블루그레이 #8995A4 계열','오른쪽 패널 안쪽 20~30%','전체 벽 사용 시 차가워질 수 있음'),
('천장','소프트 아이보리 #F3F0E9 계열','천장·코니스','순백색 대비를 너무 세게 만들지 않음'),
('금속','앤티크 브라스 소량','등기구','손잡이까지 모두 금색으로 맞추지 않음')],[1.0,1.75,2.2,1.95])
doc.add_heading('벽면 비례',2)
bullet('몰딩 폭 20~30mm, 돌출 8~12mm 정도의 얇은 프로파일을 우선합니다.')
bullet('큰 패널 2~3개로 나누고, 콘센트·문선·커튼과 몰딩 선이 충돌하지 않게 전개도를 먼저 그립니다.')
bullet('포인트 컬러는 패널 내부 일부에만 사용하고 몰딩 자체는 벽색과 같게 도장합니다.')
doc.add_heading('2. 벽면 패널 몰딩',1)
table(['항목','장점','단점','방 1실 계획금액'],[
('얇은 MDF/PU 박스 몰딩','모던 클래식 효과가 가장 분명, 컬러 분할 가능','코너 재단과 이음 퍼티 품질에 따라 완성도 차이','45~90만원'),
('PVC/스티렌 몰딩','가볍고 습기·변형 부담이 낮음, 자재비 저렴','가까이서 보면 재질감이 가벼울 수 있음','30~60만원'),
('석고 장식 몰딩','도장 일체감과 고급스러운 음영','무겁고 균열·보수 난도 높음','80~160만원')],[1.35,2.1,2.15,1.3])
doc.add_heading('권장 시공 방법',2)
for x in ['레이저로 수평·수직 기준선을 잡고 패널 전개를 확정합니다.','45도 마이터 재단 후 전용 접착제와 핀으로 고정합니다.','핀 자국과 이음부를 퍼티 처리하고 완전 건조 후 샌딩합니다.','프라이머 후 벽과 몰딩을 동일한 수성 도료로 2회 도장합니다.','사광에서 이음·울렁임·도막 뭉침을 확인합니다.']: bullet(x)
callout('조건부 추천','기존 벽지가 들뜨거나 요철이 심하면 그 위에 바로 몰딩을 붙이지 않습니다. 벽지 제거와 면 보수가 선행되어 금액이 20~50만원 이상 추가될 수 있습니다.')

page(); doc.add_heading('3. 도장과 벽지 비교',1)
table(['마감','디자인 효과','장점','단점','방 1실 계획금액'],[
('수성 도장','몰딩과 벽이 한 면처럼 연결','컬러 자유도·보수성·무광 질감 우수','퍼티·샌딩 품질이 그대로 드러남, 양생 필요','60~110만원'),
('합지 벽지','부드럽고 무난한 단색','비용이 낮고 벽 미세 요철을 일부 숨김','오염·습기에 약하고 몰딩 이음 표현 제한','35~55만원'),
('실크 벽지','질감과 오염 관리가 비교적 좋음','선택 폭 넓고 표면 내구성 양호','이음선·모서리 들뜸, 도장 같은 일체감은 약함','45~80만원'),
('프리미엄 벽지','패브릭·미장 같은 깊은 질감','포인트 벽 완성도 높음','롤 손실과 숙련 시공비가 큼','70~120만원')],[1.0,1.45,1.7,1.7,1.05])
doc.add_heading('도장 시공 핵심',2)
bullet('벽지 제거 → 접착제 제거 → 균열 보수 → 전체 퍼티 → 2회 이상 샌딩 → 프라이머 → 상도 2회가 기본입니다.')
bullet('몰딩을 먼저 설치하고 이음부를 잡은 뒤 벽과 함께 도장해야 가장 자연스럽습니다.')
bullet('무광은 고급스럽지만 손때가 남기 쉽고, 에그쉘은 은은한 광택과 닦임성이 균형 잡힙니다.')
doc.add_heading('벽지 시공 핵심',2)
bullet('몰딩 안쪽에 벽지를 넣는다면 패널별 재단선이 늘어 인건비와 들뜸 위험이 커집니다.')
bullet('모던 클래식 몰딩을 살릴 때는 전체 도장을 1순위, 벽지는 몰딩 없는 면 또는 한 면 포인트로 권합니다.')

page(); doc.add_heading('4. 바닥 걸레받이와 상부 몰딩',1)
doc.add_heading('바닥 걸레받이를 없애고 싶을 때',2)
table(['해결법','장점','단점·하자 위험','방 1실 계획금액'],[
('기존 철거 + 얇은 평걸레받이','무몰딩처럼 가볍고 벽 하단 보호 유지','완전 무몰딩은 아님','30~60만원'),
('철거 + 벽면 미장 후 무걸레받이','선이 사라져 가장 미니멀','물걸레·청소기 충격, 벽지 하단 오염, 바닥 틈 노출','60~120만원'),
('마이너스/매립 걸레받이','음영선이 정교하고 벽 보호 가능','벽체 가공·보강·정밀 재단 필요, 부분공사 비효율','90~180만원')],[1.55,1.7,2.35,1.3])
callout('이 방의 판단','모던 클래식에는 완전 무걸레받이보다 60~80mm 평걸레받이를 벽과 같은 색으로 도장하는 편이 디자인과 관리 모두 안정적입니다.',PALE,NAVY)
doc.add_heading('상부 몰딩을 없애거나 바꿀 때',2)
table(['선택','효과','주의점','방 1실 계획금액'],[
('슬림 코니스 45~60mm 교체','클래식 비례와 그림자 형성','에어컨·커튼박스 접점 상세 필요','25~55만원'),
('철거 후 무몰딩 도장','천장과 벽이 깨끗하게 연결','철거 자국·천장 수평 오차가 드러남','40~90만원'),
('마이너스 몰딩','떠 있는 천장 같은 음영선','목공·석고·도장 동반, 부분공사 비용 큼','80~180만원')],[1.5,1.8,2.35,1.25])

page(); doc.add_heading('5. 등기구와 블라인드·커튼',1)
doc.add_heading('등기구 교체',2)
table(['선택','디자인 효과','계획금액','확인사항'],[
('오팔 유리 직부/반직부','모던 클래식의 중심 포인트','제품 10~40만원 + 교체 5~15만원','무게·천장 보강·색온도'),
('패브릭 드럼 직부','빛이 부드럽고 침실에 안정적','제품 8~25만원 + 교체 5~15만원','먼지 관리·천장고'),
('간접조명 추가','몰딩 음영과 포인트색 강조','배선 포함 20~60만원 추가','스위치 회로·점검성'),
('위치 이동/천장 보수','배치 중심을 정확히 맞춤','20~50만원 추가','배선·도배/도장 보수')],[1.45,2.05,1.65,1.75])
callout('전기 안전','등기구 본체 교체라도 배선 연결과 천장 보강은 전기 전문가 작업으로 구분합니다.')
doc.add_heading('창 마감 교체',2)
table(['제품','분위기·기능','단점','창 1개 계획금액'],[
('콤비/롤 블라인드','간결하고 광량 조절 쉬움','패브릭의 부드러움은 적음','12~25만원'),
('우드 블라인드','클래식과 자연 소재감','무겁고 먼지·뒤틀림 관리','20~40만원'),
('쉬어 커튼 1겹','빛이 부드럽고 공간이 넓어 보임','야간 프라이버시·암막 부족','20~35만원'),
('쉬어+암막 이중커튼','모던 클래식 완성도·숙면·보온','부피와 비용, 세탁 필요','30~60만원')],[1.35,2.1,2.05,1.4])
bullet('천장형 에어컨 토출구와 커튼 상단이 겹치지 않도록 레일 위치와 커튼 높이를 실측합니다.')

page(); doc.add_heading('6. 디자인별 예산 시나리오',1)
table(['안','포함 범위','공사비 계획','추천 상황'],[
('A. 컬러 중심','벽지 또는 부분도장, 등기구, 창 마감','약 100~220만원','몰딩 없이 색과 빛만 바꿀 때'),
('B. 모던 클래식 균형형','패널 몰딩, 전체 도장, 슬림 코니스, 등기구, 이중커튼','약 230~430만원','이번 시안의 권장 수준'),
('C. 경계 디테일형','B안 + 걸레받이 철거/매립 + 상부 무몰딩/마이너스','약 380~700만원','면 보수와 목공까지 감수할 때')],[1.45,3.1,1.35,1.0])
doc.add_heading('견적에서 반드시 분리할 항목',2)
for x in ['기존 몰딩·벽지 철거와 폐기물','벽면 전체 퍼티와 샌딩 횟수','패널 몰딩 재질·폭·총 길이','프라이머와 상도 제품·도장 횟수','천장 코니스와 걸레받이 별도 길이','등기구 제품비·전기 인건비·천장 보강','커튼 원단·주름배수·레일·설치비','보양·청소·부가세·주차비']: bullet(x)
callout('금액 해석','표의 금액은 방 1실 사진 기반 계획 범위입니다. 최소 출장비, 지역, 벽 상태, 실제 둘레, 자재 등급에 따라 달라지므로 실측 후 공정별 수량으로 다시 산출합니다.')

page(); doc.add_heading('7. 권장 실행안',1)
callout('최종 추천','얇은 박스 몰딩 + 웜그레이지 에그쉘 도장 + 블루그레이 패널 포인트 + 45~60mm 슬림 코니스 + 벽색 도장 평걸레받이 + 오팔 반직부등 + 쉬어/암막 이중커튼',PALE,NAVY)
doc.add_heading('시공 순서',2)
for x in ['실측과 벽 전개도 작성','등기구·커튼 레일·에어컨 간섭 확인','기존 벽지 및 불량 몰딩 철거','전기 선행 작업','벽체 보수와 몰딩 목공','전체 퍼티·샌딩·프라이머','천장·벽·몰딩 도장','걸레받이와 커튼·등기구 마감','사광 검수와 보수']: bullet(x)
doc.add_heading('실측 후 확정할 디자인 값',2)
table(['확정 항목','필요 값'],[
('벽 전개','벽별 폭·높이, 문·창·콘센트 위치'),('몰딩','프로파일 샘플, 폭·돌출·패널 간격·총 길이'),('컬러','현장 A3 이상 샘플 도장, 주·야간 확인'),('경계','상부 코니스/무몰딩, 하부 평걸레받이/매립 선택'),('조명','등기구 무게·직경·색온도·조도·회로'),('창 마감','창 폭·높이·커튼박스·에어컨 토출 간격')],[1.55,5.35])

for p in doc.paragraphs: p.paragraph_format.widow_control=True
doc.core_properties.title='공간한쪽 방 모던 클래식 디자인 제안서'
doc.core_properties.subject='벽면 몰딩, 도장, 벽지, 걸레받이, 상부 몰딩, 조명, 커튼 디자인 및 계획견적'
doc.core_properties.author='공간한쪽'
doc.save(OUT); print(OUT)
