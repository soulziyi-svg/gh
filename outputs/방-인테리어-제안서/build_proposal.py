from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from pathlib import Path

BASE = Path(r"C:\ai_web\lecture1\인테리어사이트\outputs\방-인테리어-제안서")
ORIGINAL = BASE / "01_현장사진.png"
CONCEPT = BASE / "02_완성예상시안.png"
OUTPUT = BASE / "공간한쪽_방_인테리어_1차제안서.docx"

BLUE = "2864B4"
DARK = "173B67"
INK = "243447"
MUTED = "667085"
PALE = "EAF2FC"
WARM = "F6F1E9"
LINE = "D8E1EC"
RISK = "A23A3A"
GOLD = "8A671A"
FONT = "Malgun Gothic"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.7)
sec.bottom_margin = Inches(0.7)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run

styles = doc.styles
normal = styles['Normal']
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.28

for name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 18, 9),
    ('Heading 2', 13, BLUE, 12, 6),
    ('Heading 3', 11.5, DARK, 8, 4),
]:
    st = styles[name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for style_name in ['List Bullet', 'List Number']:
    st = styles[style_name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    st.font.size = Pt(10.2)
    st.paragraph_format.left_indent = Inches(0.38)
    st.paragraph_format.first_line_indent = Inches(-0.19)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.2

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn('w:' + m))
        if node is None:
            node = OxmlElement('w:' + m)
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')

def set_table_geometry(table, widths_in):
    total = int(sum(widths_in) * 1440)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:w'), str(total)); tblW.set(qn('w:type'), 'dxa')
    tblInd = tblPr.find(qn('w:tblInd'))
    if tblInd is None:
        tblInd = OxmlElement('w:tblInd'); tblPr.append(tblInd)
    tblInd.set(qn('w:w'), '120'); tblInd.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths_in:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(width*1440))); grid.append(gc)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            w = int(widths_in[idx] * 1440)
            cell.width = Inches(widths_in[idx])
            tcW = cell._tc.get_or_add_tcPr().find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW'); cell._tc.get_or_add_tcPr().append(tcW)
            tcW.set(qn('w:w'), str(w)); tcW.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)

def format_table(table, header=True):
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if header and r == 0: set_cell_shading(cell, PALE)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs: set_font(run, 9.2, bold=(header and r == 0), color=(DARK if header and r == 0 else INK))

def add_table(headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers): table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row): cells[i].text = str(val)
    set_table_geometry(table, widths)
    format_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    set_font(p.add_run(text), 10.2)
    return p

def add_num(text):
    p = doc.add_paragraph(style='List Number')
    set_font(p.add_run(text), 10.2)
    return p

def add_callout(label, text, fill=PALE, color=DARK):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'; set_table_geometry(t, [6.9])
    c = t.cell(0, 0); set_cell_shading(c, fill)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_font(p.add_run(label + "  "), 10.5, True, color)
    set_font(p.add_run(text), 10.5, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(text), 8.5, False, MUTED, True)

def page_break():
    doc.add_page_break()

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(paragraph.add_run('공간한쪽 | 사진 기반 1차 제안서   '), 8.5, False, MUTED)
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run('SPACE HANCHOK · ROOM PROPOSAL'), 8.5, True, BLUE)
add_page_number(sec.footer.paragraphs[0])

# Cover
p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
set_font(p.add_run('공간한쪽'), 11, True, BLUE)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
set_font(p.add_run('작은 방 인테리어 1차 제안서'), 25, True, DARK)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(12)
set_font(p.add_run('사진 한 장에서 시작하는 현실적인 침실 + 작업공간 계획'), 12.5, False, MUTED)
doc.add_picture(str(CONCEPT), width=Inches(6.9))
add_caption('완성 예상 시안 · 기존 구조와 설비를 유지한 가구 배치 중심 제안')
add_callout('핵심 방향', '철거보다 유지, 유행보다 관리, 사진보다 먼저 동선과 에어컨 바람길을 확인합니다.', WARM, GOLD)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(p.add_run('작성일  2026. 09. 02  |  견적 신뢰도 L1(사진 기반)'), 9, False, MUTED)

page_break()
doc.add_heading('1. 먼저 내린 결론', level=1)
add_callout('추천안', '현재 마감은 철거하지 않고, 1200×2000mm 이하 침대와 창 아래 1000~1200mm 책상을 배치하는 저공사형 계획을 우선합니다.')
p = doc.add_paragraph('이 방은 기존 바닥·벽·문 상태가 비교적 정돈되어 보이고, 천장형 에어컨과 창 위치가 명확합니다. 큰 공사보다 가구의 크기, 수납량, 조명 색온도를 정확히 잡는 편이 비용 대비 효과가 큽니다.')
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
doc.add_heading('현장 사진', level=2)
doc.add_picture(str(ORIGINAL), width=Inches(6.9))
add_caption('제공 사진 · 실제 치수와 벽체 상태는 현장 실측 전 추정')

page_break()
doc.add_heading('2. 현장 조건 분석', level=1)
add_table(['구분', '사진에서 확인되는 내용', '설계 영향'], [
    ('확인', '긴 직사각형 방, 왼쪽 출입문, 정면 창', '문에서 창까지의 중앙 통로를 우선 확보'),
    ('확인', '창 상부 천장형 에어컨', '창 앞 높은 가구와 침대 머리 위치를 피함'),
    ('확인', '밝은 바닥·벽·문, 걸레받이 유지 상태 양호', '철거 없이 보수·세척 중심으로 비용 절감'),
    ('확인', '왼쪽 벽 하부 콘센트 1개 노출', '책상 전원과 침대 충전 동선 추가 확인'),
    ('추정', '방 폭 약 2.8~3.2m, 길이 약 3.8~4.4m', '1200mm 침대 적용 가능성은 높지만 실측 필수'),
    ('미확인', '벽체 종류, 누수·곰팡이, 전기 용량, 반입 경로', '시공 전 체크리스트에서 확정'),
], [0.85, 3.1, 2.95])
doc.add_heading('실측 때 반드시 받을 8개 치수', level=2)
for item in ['방 가로·세로·천장고', '문 폭·문 열림 반경', '창 전체 폭·높이·바닥에서 창턱 높이', '에어컨 토출구 위치와 필터 점검 여유', '모든 콘센트·스위치 위치', '걸레받이 돌출 깊이', '엘리베이터·복도·현관 반입 폭', '침대 옆 최소 통로 폭']:
    add_bullet(item)
add_callout('중요', '시안 속 가구는 비례 제안입니다. 실측 후에는 침대 옆 통로 800mm 전후, 책상 깊이 500~600mm, 의자 뒤 사용 여유 700mm 이상을 다시 검증합니다.', WARM, GOLD)

page_break()
doc.add_heading('3. 설계 시안과 배치 원칙', level=1)
doc.add_picture(str(CONCEPT), width=Inches(5.5))
add_caption('따뜻한 미니멀 침실 + 집중 가능한 작은 작업공간')
add_table(['영역', '제안 규격(실측 전)', '배치 이유'], [
    ('침대', '슈퍼싱글~작은 더블, W1100~1200 × L2000', '오른쪽 긴 벽을 사용해 출입 동선 분리'),
    ('책상', 'W1000~1200 × D500~600', '창 아래 낮은 가구로 채광 활용, 블라인드 조작 유지'),
    ('수납', 'W800~1000 × D300~400 저상장', '문 간섭을 피하고 시야를 막지 않음'),
    ('조명', '주조명 3500~4000K + 침대 간접 3000K', '작업성과 휴식성을 분리'),
], [1.0, 2.5, 3.4])
doc.add_heading('디자인 언어', level=2)
add_bullet('벽과 바닥은 유지하고 웜화이트·오트밀·내추럴 오크를 중심으로 정리합니다.')
add_bullet('포인트 컬러는 사이트 방향과 연결되는 낮은 채도의 블루그레이를 5% 이내로 사용합니다.')
add_bullet('장식보다 청소가 쉬운 닫힌 수납, 세탁 가능한 침구, 이동 가능한 기성가구를 우선합니다.')
doc.add_heading('4. 선택의 장단점 비교', level=1)
add_table(['선택', '장점', '주의점', '판정'], [
    ('기존 벽지 유지 + 부분보수', '비용·먼지·공기 단축', '색 차이와 찍힘이 남을 수 있음', '조건부 추천'),
    ('전체 도배', '면이 균일하고 새집 느낌', '가구 반입 전 공정 필요, 폐기물 발생', '오염이 넓을 때'),
    ('셀프 수성도장', '색 선택 자유, 부분 재도장 가능', '벽지 이음·들뜸 보수와 양생이 어려움', '작은 면만 추천'),
    ('맞춤 붙박이장', '수납 밀도와 마감 우수', '가격·철거·이사 대응 불리', '실측 후 판단'),
    ('기성 저상 수납', '저렴하고 이동·교체 쉬움', '빈틈과 수납 효율은 낮음', '1차 추천'),
    ('현재 암막 블라인드 유지', '추가비 없음, 암막성', '어두운 색이 방을 좁게 보이게 함', '상태 양호 시 유지'),
], [1.45, 1.8, 2.25, 1.4])
doc.add_heading('이 방에서 하지 않을 것', level=2)
add_bullet('사진만 보고 무몰딩·마이너스몰딩 공사를 권하지 않습니다. 천장형 에어컨과 기존 마감 연결부까지 손대면 비용과 하자 위험이 커집니다.')
add_bullet('창 앞에 높은 장이나 침대 머리를 두지 않습니다. 냉난방 토출과 필터 점검, 블라인드 조작을 방해할 수 있습니다.')
add_bullet('예쁜 사진을 위해 통로를 600mm 이하로 줄이지 않습니다. 매일 쓰는 공간은 촬영 구도보다 이동과 청소가 먼저입니다.')

page_break()
doc.add_heading('5. 셀프와 전문가 작업 구분', level=1)
add_table(['작업', '셀프 가능 조건', '전문가를 부를 기준'], [
    ('세척·가구 배치·침구', '가능. 바닥 보양 후 조립', '대형 가구 반입·벽 고정이 필요할 때'),
    ('부분 벽지 보수', '작은 찢김·오염, 동일 자재 확보', '들뜸이 넓거나 곰팡이·누수 흔적이 있을 때'),
    ('수성도장', '작은 포인트 면, 충분한 환기·양생', '전체 벽, 벽지 이음 보수, 색 균일성이 중요할 때'),
    ('조명 교체·콘센트 증설', '등기구 외관 청소만', '전기 연결·배선은 전기 전문가'),
    ('벽 선반·헤드보드 고정', '벽체 확인과 적정 앵커 사용 가능', '석고보드 보강·중량물·배관 위치 불명'),
], [1.35, 2.65, 2.9])
doc.add_heading('권장 작업 순서', level=2)
for step in ['실측·사진 추가 촬영', '누수·곰팡이·전기·에어컨 점검', '벽 보수 또는 도배 판단', '조명 전기 작업', '청소와 바닥 보양', '큰 가구부터 반입·수평 조정', '블라인드·조명·침구 마감', '통로·문·에어컨 작동 최종 확인']:
    add_num(step)

page_break()
doc.add_heading('6. 1차 예산과 구매 기준', level=1)
add_callout('예산 기준', '아래 금액은 2026-09-02 사진 기반 계획예산입니다. 배송·조립·지역 인건비·기존 하자에 따라 달라지며, 발주 전 판매가와 현장 견적을 다시 확인해야 합니다.', WARM, GOLD)
add_table(['항목', '수량/기준', '계획 범위', '비고'], [
    ('벽 세척·부분보수', '방 1실', '10~30만원', '셀프 자재 또는 소규모 출장'),
    ('전체 도배/도장 선택 시', '벽+천장', '50~100만원', '현장 상태·지역별 변동'),
    ('조명 보완', '주조명 1 + 간접 2', '15~35만원', '전기 작업 포함 여부 확인'),
    ('침대 프레임', '120×200급 1', '11.9~39.4만원', '매트리스 별도'),
    ('매트리스', '120×200급 1', '25~55만원', '누워보고 선택 권장'),
    ('책상', '100~120cm 1', '8~18.9만원', '깊이와 창턱 간섭 확인'),
    ('의자', '1', '8~25만원', '좌판 높이·등받이 우선'),
    ('저상 수납장', '1', '15~40만원', '문 열림 반경 확인'),
    ('블라인드 선택 교체', '창 1', '15~35만원', '실측 제작'),
    ('배송·조립·예비비', '합계의 약 10~15%', '20~50만원', '반입·폐기 포함 확인'),
], [1.55, 1.35, 1.25, 2.75])
doc.add_heading('총액 시나리오', level=2)
add_table(['안', '범위', '포함 기준'], [
    ('유지형', '약 110~210만원', '기존 마감 유지, 기성가구, 셀프 조립 중심'),
    ('균형형', '약 190~350만원', '부분 마감 보수, 조명·블라인드, 중급 가구와 조립'),
    ('맞춤형', '350만원 이상', '맞춤 헤드보드·수납, 전체 마감, 전문 시공 확대'),
], [1.2, 1.6, 4.1])

page_break()
doc.add_heading('7. 구매 후보와 발주 체크', level=1)
add_table(['품목', '현재 확인 예시', '선정 기준'], [
    ('침대 프레임', 'IKEA VEVELSTAD 120×200, 119,000원', '방 폭 실측 후 통로 800mm 전후 확보'),
    ('침대 대안', 'IKEA MALM 120×200 오크 무늬, 394,000원', '따뜻한 톤은 좋지만 프레임 외곽 치수 확인'),
    ('매트리스 예시', 'IKEA VALEVÅG 120×200, 319,000원', '체형·경도는 매장 체험 후 결정'),
    ('책상 예시', 'IKEA MITTZON 120×60, 169,000~189,000원', '창턱·블라인드 체인·의자 후퇴 공간 확인'),
    ('저가 책상 구성', 'LAGKAPTEN 상판 120×60, 39,900원부터', '다리·서랍·배송비는 별도 합산'),
], [1.35, 3.2, 2.35])
p = doc.add_paragraph('구매 링크 및 확인일')
p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(4)
set_font(p.runs[0], 9.2, True, DARK)
for url in [
    'https://www.ikea.com/kr/ko/cat/beds-bm003/',
    'https://www.ikea.com/kr/ko/p/malm-bed-frame-high-white-stained-oak-veneer-loenset-s69157296/',
    'https://www.ikea.com/kr/ko/p/mittzon-desk-white-s59525845/',
    'https://www.ikea.com/kr/ko/cat/table-tops-11844/',
]:
    p = doc.add_paragraph(style='List Bullet'); set_font(p.add_run(url), 8.6, False, MUTED)
add_callout('발주 전 멈춤 지점', '침대·책상·수납장을 한꺼번에 주문하지 않습니다. 실측표에 외곽 치수를 적고, 문 열림과 의자 사용 상태를 바닥 마스킹테이프로 먼저 표시한 뒤 발주합니다.', PALE, DARK)

page_break()
doc.add_heading('8. 다음 단계: 현장 확정 프로토콜', level=1)
for title, desc in [
    ('01 입력', '정면·좌·우·출입문 사진, 방 치수, 창·문·콘센트 치수, 예산, 보유 가구를 받습니다.'),
    ('02 조건 분석', '동선, 채광, 냉난방, 벽체, 누수, 전기, 반입·폐기 조건을 확인합니다.'),
    ('03 시안 확정', '가구 외곽 치수와 통로를 2D 배치로 검증하고, 색·재료를 결정합니다.'),
    ('04 비교', '유지/교체, 셀프/전문가, 초기비/관리비의 장단점을 표로 비교합니다.'),
    ('05 견적', '수량 산출, 단가 출처, 배송·조립·폐기·예비비를 분리해 확정합니다.'),
    ('06 실행', '선행 공정부터 발주하고, 각 단계 사진으로 검수한 뒤 다음 공정을 진행합니다.'),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    set_font(p.add_run(title + '  '), 11, True, BLUE)
    set_font(p.add_run(desc), 10.5, False, INK)
doc.add_heading('현장 확정 전 체크박스', level=2)
for item in ['침대 외곽 치수를 바닥에 표시했다', '문이 끝까지 열리고 통로가 확보된다', '책상 의자를 뺀 상태에서도 이동 가능하다', '에어컨 토출·필터 점검을 막지 않는다', '콘센트 멀티탭이 통로를 가로지르지 않는다', '벽 고정 전 배관·배선·벽체를 확인했다', '배송·조립·포장재 회수 비용을 확인했다']:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run('□  ' + item), 10.3)
add_callout('제안서의 현재 상태', '사진 기반 1차안입니다. 실측표와 추가 사진이 들어오면 배치 치수, 자재 수량, 공정별 견적, 구매 링크를 L2 견적으로 갱신할 수 있습니다.', WARM, GOLD)

# Keep headings with following content and prevent accidental widows where possible.
for p in doc.paragraphs:
    if p.style.name.startswith('Heading'):
        p.paragraph_format.keep_with_next = True
    p.paragraph_format.widow_control = True

doc.core_properties.title = '공간한쪽 작은 방 인테리어 1차 제안서'
doc.core_properties.subject = '사진 기반 현장 분석, 설계 시안, 장단점, 셀프/전문가 구분, 계획 예산'
doc.core_properties.author = '공간한쪽'
doc.save(OUTPUT)
print(OUTPUT)
